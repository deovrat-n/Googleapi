from docx import Document
from docx.shared import Inches

# Create a new document
document = Document()

# Add a heading with the name
document.add_heading('DEOVRAT NETAM', 0)

# Add a paragraph with the enrollment number
document.add_paragraph('Enrollment No.: 191010214')

# Add a paragraph with the contact information
document.add_paragraph('Phone: +91-8770438596')
document.add_paragraph('Email: deovrat19101@iiitnr.edu.in')

# Add a paragraph with the social media links
document.add_paragraph('GitHub: https://github.com/deovrat-n')
document.add_paragraph('LinkedIn: https://www.linkedin.com/in/deovrat-netam-8a876420b/')

# Add a heading for the education section
document.add_heading('EDUCATION', 1)

# Add a paragraph for the B.Tech degree
document.add_paragraph('B.Tech - Electronics and Communication Engineering, International Institute of Information Technology, Naya Raipur (IIIT-NR), 2019-Present, CGPA: 7.5')

# Add a paragraph for the Senior Secondary education
document.add_paragraph('Senior Secondary, CBSE, 2019, 71.4%')

# Add a paragraph for the Secondary education
document.add_paragraph('Secondary, CBSE, 2017, 10')

# Add a heading for the technical projects section
document.add_heading('TECHNICAL PROJECTS', 1)

# Add a paragraph for the Stock Price Prediction project
document.add_paragraph('**Stock Price Prediction:** Forecast stock prices using LSTM (Long Short-Term Memory) Model, achieving high accuracy in predicting Infosys\' stock price.')

# Add a paragraph for the Educational Game project
document.add_paragraph('**Educational Game:** Developed a game in Unity Engine to teach Real Numbers, utilizing C#, object-oriented programming, and Visual Studio.')

# Add a paragraph for the Text-based Adventure Game project
document.add_paragraph('**Text-based Adventure Game:** Created a game in C# where players navigate a dungeon, using Unity Engine, C#, object-oriented programming, and Data Structures.')

# Add a paragraph for the Text Summarizer project
document.add_paragraph('**Text Summarizer:** Implemented an encoder-decoder architecture (using LSTM) to create text summaries, employing an attention method to retain information.')

# Add a heading for the positions of responsibility section
document.add_heading('POSITIONS OF RESPONSIBILITY', 1)

# Add a paragraph for the Game Development Challenge
document.add_paragraph('Led a team in Game Development Challenge organized by CDAC-Mumbai and IEEE Bombay Section')

# Add a paragraph for the Football Club
document.add_paragraph('Core Member of Football Club, IIIT Naya Raipur')

# Add a paragraph for the Vriddhi'19 - NIT Rourkela Sports Fest
document.add_paragraph('Represented IIIT Naya Raipur at Vriddhi\'19 - NIT Rourkela Sports Fest')

# Add a heading for the achievements section
document.add_heading('ACHIEVEMENTS', 1)

# Add a paragraph for the EG-100 Educational Game Challenge
document.add_paragraph('First Consolation Prize in EG-100 Educational Game Challenge organized by CDAC-Mumbai and IEEE Bombay Section')

# Add a heading for the technical skills section
document.add_heading('TECHNICAL SKILLS', 1)

# Add a paragraph for the languages
document.add_paragraph('Languages: HTML, CSS, C, C++, Python, Java, C#, SQL')

# Add a paragraph for the frameworks and libraries
document.add_paragraph('Frameworks and Libraries: Visual Studio; Microsoft SQL Server; Scikit-learn; PyTorch')

# Add a heading for the soft skills section
document.add_heading('SOFT SKILLS', 1)

# Add a paragraph for the leadership skill
document.add_paragraph('Leadership')

# Add a paragraph for the creativity skill
document.add_paragraph('Creativity')

# Add a paragraph for the teamwork skill
document.add_paragraph('Teamwork')

# Save the document
document.save('resume.docx')